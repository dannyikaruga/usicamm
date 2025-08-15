import os
import re
import pytesseract
import fitz  # PyMuPDF
from PIL import Image
import tempfile
import subprocess
from docx import Document
from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename

# Configuración Flask
app = Flask(__name__)
UPLOAD_FOLDER = r"D:\Ussicamm\AI\pdfs"
OUTPUT_FOLDER = r"D:\Ussicamm\AI\results_word"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Ajusta Tesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
os.environ["TESSDATA_PREFIX"] = r"C:\Program Files\Tesseract-OCR\tessdata"

MAX_CHARS = 2000

def clean_text(text):
    text = re.sub(r'[^\x20-\x7EÁÉÍÓÚáéíóúÜüÑñ.,;:!?()\-–—\n]', '', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'(\S{50,})', lambda m: m.group(1)[:50] + '-', text)
    return text.strip()

def extract_text_from_pdf(pdf_path):
    logs = []
    logs.append("🔹 Abriendo PDF para extracción de texto...")
    text = ""
    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)
    logs.append(f"📄 Total de páginas: {total_pages}")

    for page_num in range(total_pages):
        logs.append(f"➡ Procesando página {page_num + 1}/{total_pages}...")
        page = pdf_doc.load_page(page_num)

        page_text = page.get_text()
        if page_text.strip():
            text += page_text + "\n"
        else:
            pix = page.get_pixmap(dpi=300)
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                tmp_img.write(pix.tobytes())
                tmp_img_path = tmp_img.name

            img = Image.open(tmp_img_path).convert("L")
            img = img.point(lambda x: 0 if x < 140 else 255, '1')
            custom_config = r'--psm 6'
            ocr_text = pytesseract.image_to_string(img, lang="spa+eng", config=custom_config)
            text += ocr_text + "\n"
            os.remove(tmp_img_path)

    pdf_doc.close()
    text = clean_text(text)
    logs.append(f"✅ Extracción completa, {len(text)} caracteres obtenidos.")
    return text, logs

def split_into_blocks(text, max_chars=MAX_CHARS):
    return [text[i:i+max_chars] for i in range(0, len(text), max_chars)]

def query_ollama(prompt, model="deepseek-llm:7b"):
    process = subprocess.Popen(
        ["ollama", "run", model],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.DEVNULL
    )
    output_bytes, _ = process.communicate(input=prompt.encode('utf-8'))
    return output_bytes.decode('utf-8', errors='ignore').strip()

def save_word_and_convert_to_pdf(pdf_name, analysis_results):
    doc = Document()
    doc.add_heading(f"Requisitos extraídos de la convocatoria: {pdf_name}", level=1)
    doc.add_paragraph(analysis_results)
    word_path = os.path.join(OUTPUT_FOLDER, pdf_name.replace(".pdf", ".docx"))
    doc.save(word_path)

    subprocess.run([
        r"C:\Program Files\LibreOffice\program\soffice.exe", "--headless",
        "--convert-to", "pdf", word_path, "--outdir", OUTPUT_FOLDER
    ], check=True)

@app.route("/")
def index():
    return render_template("index_doc.html")

@app.route("/upload", methods=["POST"])
def upload_pdf():
    if "pdf" not in request.files:
        return jsonify({"error": "No se envió archivo"})
    file = request.files["pdf"]
    if file.filename == "":
        return jsonify({"error": "Nombre de archivo vacío"})
    filename = secure_filename(file.filename)
    file.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
    return jsonify({"message": f"Archivo {filename} subido correctamente"})

@app.route("/process", methods=["POST"])
def process_pdfs():
    logs_total = []
    pdfs = [f for f in os.listdir(UPLOAD_FOLDER) if f.lower().endswith(".pdf")]
    if not pdfs:
        return jsonify({"logs": ["⚠ No se encontraron PDFs en la carpeta de subida."]})

    for pdf_file in pdfs:
        logs_total.append(f"🔹 Procesando {pdf_file}...")
        pdf_path = os.path.join(UPLOAD_FOLDER, pdf_file)
        text, logs = extract_text_from_pdf(pdf_path)
        logs_total.extend(logs)
        blocks = split_into_blocks(text)

        full_analysis = ""
        for i, block in enumerate(blocks, start=1):
            logs_total.append(f"🤖 Procesando bloque {i}/{len(blocks)} con Ollama...")
            prompt = (
                "Extrae únicamente los requisitos de la convocatoria en el siguiente texto. "
                "No agregues encabezados ni numeraciones:\n\n" + block
            )
            result = query_ollama(prompt)
            full_analysis += result + "\n\n"

        save_word_and_convert_to_pdf(pdf_file, full_analysis)
        logs_total.append(f"✅ {pdf_file} completado.")

    return jsonify({"logs": logs_total})

if __name__ == "__main__":
    app.run(debug=True)
