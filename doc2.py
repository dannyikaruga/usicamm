import os
import re
import pytesseract
import fitz  # PyMuPDF
from PIL import Image
import tempfile
import subprocess
from docx import Document

# Ajusta la ruta a tesseract si no está en PATH
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
os.environ["TESSDATA_PREFIX"] = r"C:\Program Files\Tesseract-OCR\tessdata"

# Carpetas de entrada y salida
pdf_folder = r"D:\Ussicamm\AI\pdfs"
output_folder = r"D:\Ussicamm\AI\results_word"
os.makedirs(output_folder, exist_ok=True)

# Máximo caracteres por bloque
MAX_CHARS = 2000

def clean_text(text):
    # Mantener solo caracteres imprimibles y acentos
    text = re.sub(r'[^\x20-\x7EÁÉÍÓÚáéíóúÜüÑñ.,;:!?()\-–—\n]', '', text)
    # Reemplazar múltiples espacios
    text = re.sub(r'\s+', ' ', text)
    # Reemplazar secuencias largas de caracteres sin espacios por guiones
    text = re.sub(r'(\S{50,})', lambda m: m.group(1)[:50] + '-', text)
    return text.strip()

def extract_text_from_pdf(pdf_path):
    print("🔹 Abriendo PDF para extracción de texto...")
    text = ""
    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)
    print(f"📄 Total de páginas: {total_pages}")

    for page_num in range(total_pages):
        print(f"➡ Procesando página {page_num + 1}/{total_pages}...")
        page = pdf_doc.load_page(page_num)

        page_text = page.get_text()
        if page_text.strip():
            text += page_text + "\n"
        else:
            pix = page.get_pixmap(dpi=300)
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                tmp_img.write(pix.tobytes())
                tmp_img_path = tmp_img.name

            img = Image.open(tmp_img_path).convert("L")
            img = img.point(lambda x: 0 if x < 140 else 255, '1')
            custom_config = r'--psm 6'
            ocr_text = pytesseract.image_to_string(img, lang="spa+eng", config=custom_config)
            text += ocr_text + "\n"

            os.remove(tmp_img_path)

    pdf_doc.close()
    text = clean_text(text)
    print(f"✅ Extracción completa, {len(text)} caracteres obtenidos.")
    return text

def split_into_blocks(text, max_chars=MAX_CHARS):
    blocks = [text[i:i+max_chars] for i in range(0, len(text), max_chars)]
    print(f"✂ Texto dividido en {len(blocks)} bloque(s) de máximo {max_chars} caracteres.")
    return blocks

def query_ollama(prompt, model="deepseek-llm:7b"):
    process = subprocess.Popen(
        ["ollama", "run", model],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.DEVNULL
    )
    output_bytes, _ = process.communicate(input=prompt.encode('utf-8'))
    output = output_bytes.decode('utf-8', errors='ignore')
    return output.strip()

def save_word_and_convert_to_pdf(pdf_name, analysis_results):
    # Guardar en Word
    doc = Document()
    doc.add_heading(f"Requisitos extraídos de la convocatoria: {pdf_name}", level=1)
    doc.add_paragraph(analysis_results)
    word_path = os.path.join(output_folder, pdf_name.replace(".pdf", ".docx"))
    doc.save(word_path)
    print(f"💾 Resultado guardado en Word: {word_path}")

    # Convertir a PDF usando LibreOffice headless
    pdf_path = os.path.join(output_folder, pdf_name.replace(".pdf", ".pdf"))
    try:
        # Ajusta la ruta a soffice.exe si no está en PATH
        subprocess.run([
            r"C:\Program Files\LibreOffice\program\soffice.exe", "--headless",
            "--convert-to", "pdf", word_path, "--outdir", output_folder
        ], check=True)
        print(f"💾 Resultado convertido a PDF: {pdf_path}")
    except subprocess.CalledProcessError as e:
        print(f"❌ Error al convertir a PDF con LibreOffice: {e}")

def main():
    pdfs = [f for f in os.listdir(pdf_folder) if f.lower().endswith(".pdf")]
    if not pdfs:
        print(f"⚠ No se encontraron archivos PDF en la carpeta: {pdf_folder}")
        return

    print(f"📂 Se encontraron {len(pdfs)} archivo(s) PDF para procesar.")

    for idx, pdf_file in enumerate(pdfs, start=1):
        print(f"\n🔹 [{idx}/{len(pdfs)}] Procesando archivo: {pdf_file}")
        pdf_path = os.path.join(pdf_folder, pdf_file)
        text = extract_text_from_pdf(pdf_path)
        blocks = split_into_blocks(text)

        full_analysis = ""
        total_blocks = len(blocks)
        for i, block in enumerate(blocks, start=1):
            print(f"🤖 Procesando bloque {i}/{total_blocks} con Ollama...")
            prompt = (
                "Extrae únicamente los requisitos de la convocatoria en el siguiente texto. "
                "No agregues encabezados, numeraciones, ni menciones a bloques. "
                "Devuélvelo en un listado limpio y claro:\n\n"
                + block
            )
            result = query_ollama(prompt)
            print(f"✅ Bloque {i} procesado, {len(result)} caracteres obtenidos.")
            full_analysis += result + "\n\n"

        save_word_and_convert_to_pdf(pdf_file, full_analysis)
        print(f"🎯 Archivo {pdf_file} completado.")

    print("\n🏁 Todos los archivos PDF han sido procesados correctamente.")

if __name__ == "__main__":
    main()
