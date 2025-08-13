import os
import re
import csv
import pytesseract
import fitz  # PyMuPDF
from PIL import Image
import tempfile
import subprocess
from docx import Document

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

pdf_folder = r"D:\Ussicamm\AI\pdfs"
output_folder = r"D:\Ussicamm\AI\results_word"
os.makedirs(output_folder, exist_ok=True)

MAX_CHARS = 2000
RESPONSES_CSV = r"D:\Ussicamm\AI\responses.csv"

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x00-\x7F\u00C0-\u017F]+', ' ', text)
    return text.strip()

def extract_text_from_pdf(pdf_path):
    text = ""
    pdf_doc = fitz.open(pdf_path)

    for page_num in range(len(pdf_doc)):
        page = pdf_doc.load_page(page_num)

        page_text = page.get_text()
        if page_text.strip():
            text += page_text + "\n"
        else:
            pix = page.get_pixmap()
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                tmp_img.write(pix.tobytes())
                tmp_img_path = tmp_img.name
            ocr_text = pytesseract.image_to_string(Image.open(tmp_img_path), lang="spa")
            text += ocr_text + "\n"
            os.remove(tmp_img_path)

    pdf_doc.close()
    return clean_text(text)

def split_into_blocks(text, max_chars=MAX_CHARS):
    return [text[i:i+max_chars] for i in range(0, len(text), max_chars)]

def query_ollama(prompt, model="deepseek-llm:7b"):
    process = subprocess.Popen(
        ["ollama", "run", model],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.DEVNULL
    )
    output_bytes, _ = process.communicate(input=prompt.encode('utf-8'))
    output = output_bytes.decode('utf-8', errors='ignore')
    return output.strip()

def save_results_to_word(pdf_name, analysis_results):
    doc = Document()
    doc.add_heading(f"Requisitos extraídos de la convocatoria: {pdf_name}", level=1)
    doc.add_paragraph(analysis_results)
    output_path = os.path.join(output_folder, pdf_name.replace(".pdf", ".docx"))
    doc.save(output_path)
    print(f"✅ Resultado guardado en: {output_path}")

def save_response_csv(question, answer):
    try:
        exists = os.path.exists(RESPONSES_CSV)
        with open(RESPONSES_CSV, "a", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            if not exists:
                writer.writerow(["question", "answer"])
            writer.writerow([question, answer])
        print(f"✔ Guardado Q: {question[:30]}...")
    except Exception as e:
        print(f"❌ Error guardando CSV: {e}")

def parse_faq_response(text):
    if "||" in text:
        parts = text.split("||", 1)
        return parts[0].strip(), parts[1].strip()
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if len(lines) >= 2:
        return lines[0], " ".join(lines[1:])
    return None, None

def main():
    for pdf_file in os.listdir(pdf_folder):
        if not pdf_file.lower().endswith(".pdf"):
            continue
        pdf_path = os.path.join(pdf_folder, pdf_file)
        print(f"\n📄 Procesando {pdf_file}...")
        text = extract_text_from_pdf(pdf_path)
        blocks = split_into_blocks(text)

        full_analysis = ""
        for block in blocks:
            prompt = (
                "Extrae únicamente los requisitos de la convocatoria en el siguiente texto. "
                "No agregues encabezados, numeraciones, ni menciones a bloques. "
                "Devuélvelo en un listado limpio y claro:\n\n"
                + block
            )
            result = query_ollama(prompt)
            full_analysis += result + "\n\n"

            faq_prompt = (
                "Genera una pregunta frecuente relacionada a los requisitos que aparecen en este texto, "
                "y la respuesta clara y breve a esa pregunta. "
                "Devuelve solo la pregunta y respuesta separadas por '||':\n\n"
                + block
            )
            faq_result = query_ollama(faq_prompt)
            print(f"📋 FAQ generado: {faq_result}")

            q, a = parse_faq_response(faq_result)
            if q and a:
                save_response_csv(q, a)
            else:
                print("⚠ No se pudo extraer pregunta y respuesta del texto del LLM.")

        save_results_to_word(pdf_file, full_analysis)

if __name__ == "__main__":
    main()
